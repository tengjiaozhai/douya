"""PageIndexRAG 文件解析器。

这个模块负责把上传文件（txt/pdf/docx 等）统一转换为：
1) 可入库的分页文本 `pages`
2) 文档来源类型 `source_type`
3) 解析期元数据 `metadata`（OCR 摘要、图片资产等）
"""

# 启用延迟类型注解，避免前向引用在运行时造成导入顺序问题。
from __future__ import annotations

# 用于把图片字节编码成 base64，传给 Java 上传接口。
import base64
# 用于序列化/反序列化 JSON 请求与响应。
import json
# 统一日志记录，便于排查上传失败、OCR 失败等问题。
import logging
# 读取环境变量开关与阈值配置。
import os
# dataclass 用于定义不可变的解析结果对象。
from dataclasses import dataclass, field
# BytesIO 把 bytes 包装成类文件对象，供 pypdf/docx 等库读取。
from io import BytesIO
# Path 负责处理文件名后缀、stem 等路径相关逻辑。
from pathlib import Path
# Any 用于标注 OCR 返回结构等动态类型。
from typing import Any
# 捕获 HTTP 上传时的网络类错误。
from urllib.error import URLError
# 使用标准库发起 HTTP POST，把图片上传到 Java 接口。
from urllib.request import Request, urlopen

# 复用统一文本规范化与分页切分能力。
from app.indexing.chunker import normalize_text, split_pages

# 当前模块日志器，logger 名称即模块全路径。
logger = logging.getLogger(__name__)


# 业务异常：上传了当前解析器不支持的文件类型。
class UnsupportedFileTypeError(ValueError):
    # 使用 pass 保持异常类型轻量，语义通过类名表达。
    pass


# 解析后的统一文档结构。
@dataclass(frozen=True)
class ParsedDocument:
    # 入库页面文本列表；每个元素对应一个“检索页面单元”。
    pages: list[str]
    # 源类型：text/pdf/pdf+ocr/docx 等。
    source_type: str
    # 解析阶段补充的元数据（默认空对象）。
    metadata: dict[str, Any] = field(default_factory=dict)


# 视为纯文本可直接解码的后缀集合。
_TEXT_EXTS = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".tsv",
    ".json",
    ".jsonl",
    ".yaml",
    ".yml",
    ".xml",
    ".html",
    ".htm",
    ".log",
    ".ini",
    ".conf",
}


# 统一入口：根据后缀路由到不同解析分支。
def parse_uploaded_document(filename: str, data: bytes) -> ParsedDocument:
    # 空文件直接拒绝，避免后续库报错信息不清晰。
    if not data:
        raise ValueError("Uploaded file is empty.")
    # 获取文件后缀并标准化为小写，便于稳定匹配。
    ext = Path(filename).suffix.lower()

    # 文本类文件（或无后缀）走“解码 + 自动分页”路径。
    if ext in _TEXT_EXTS or not ext:
        # 多编码兜底解码。
        content = _decode_text(data)
        # 按 chunker 规则切成页面/段落片段。
        pages = split_pages(content)
        # 文本类通常没有额外解析元数据。
        return ParsedDocument(pages=pages, source_type="text", metadata={})

    # PDF 走“原生提取 + OCR 回退 + 图片资产上传”路径。
    if ext == ".pdf":
        return _parse_pdf_document(filename, data)

    # DOCX 走 docx 解析后再分页。
    if ext == ".docx":
        # 提取正文文本。
        content = _parse_docx_text(data)
        # 复用统一分页策略。
        pages = split_pages(content)
        # DOCX 当前不附带额外 metadata。
        return ParsedDocument(pages=pages, source_type="docx", metadata={})

    # 其他后缀不支持，明确告知调用方支持范围。
    raise UnsupportedFileTypeError(
        f"Unsupported file type '{ext}'. Supported: text/*, .pdf, .docx"
    )


# 文本字节解码：按常见编码优先级依次尝试。
def _decode_text(data: bytes) -> str:
    # UTF-8(BOM) -> UTF-8 -> GB18030（兼容中文 Windows 文本）。
    for enc in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            # 当前编码解码成功则直接返回。
            return data.decode(enc)
        except UnicodeDecodeError:
            # 当前编码失败则继续尝试下一种。
            continue
    # 最后兜底 latin-1，保证“尽量不抛解码异常”。
    return data.decode("latin-1")


# 纯 PDF 文本抽取（严格模式）：若无可提取文本则报错。
def _parse_pdf_pages(data: bytes) -> list[str]:
    try:
        # 运行时导入，降低模块初始加载依赖压力。
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - import error depends on environment
        # 缺依赖时给出明确安装提示。
        raise RuntimeError("pypdf is required for PDF upload. Please install pypdf.") from exc

    # 用内存流包装 bytes，交给 pypdf。
    reader = PdfReader(BytesIO(data))
    # 存放过滤空页后的文本结果。
    pages: list[str] = []
    # 逐页读取 PDF。
    for page in reader.pages:
        # 提取文本后做规范化（去噪/空白规整）。
        text = normalize_text(page.extract_text() or "")
        # 非空文本才保留。
        if text:
            pages.append(text)
    # 所有页都提取不到文本，判定为不可提取。
    if not pages:
        raise ValueError("PDF contains no extractable text.")
    # 返回文本页列表。
    return pages


# PDF 主流程：原生文本 + OCR 回退 + 可选 OSS 图片上传。
def _parse_pdf_document(filename: str, data: bytes) -> ParsedDocument:
    # 允许空页结果（空字符串），后续按页判断是否触发 OCR。
    pages_from_pdf = _extract_pdf_pages_allow_empty(data)
    # 生成安全 stem（用于图片资产命名）。
    doc_stem = _sanitize_file_stem(filename)
    # 最终用于入库的页面内容列表。
    result_pages: list[str] = []
    # 记录上传成功的页图资产元数据。
    page_image_assets: list[dict[str, Any]] = []

    # 记录实际使用 OCR 的页数。
    ocr_pages = 0
    # 记录原生文本可提取页数。
    native_pages = 0
    # 保存 OCR 过程中最后一次错误，便于最终报错上下文。
    ocr_error: str | None = None
    # 小于该字符阈值的页面视为“文本不足”，触发 OCR 回退。
    min_chars = _env_int("PAGE_INDEX_RAG_OCR_MIN_TEXT_CHARS", 20)

    # 遍历 PDF 每页；page_idx 从 1 开始便于人类阅读和展示。
    for page_idx, native_text in enumerate(pages_from_pdf, start=1):
        # 原生文本非空时累计计数。
        if native_text:
            native_pages += 1
        # 每页单独维护渲染得到的 PNG 字节（默认无）。
        page_png = None
        # 每页 OCR 文本初始化为空。
        ocr_text = ""

        # 当原生文本长度过短时，通常说明扫描件/图片页，走 OCR 回退。
        if len(native_text) < min_chars:
            # 注意 fitz 页索引从 0 开始，因此传 page_idx - 1。
            page_png, render_error = _render_pdf_page_to_png(data, page_idx - 1)
            # 记录渲染失败原因（若有）。
            if render_error:
                ocr_error = render_error
            # 渲染成功才进入 OCR。
            if page_png:
                # OCR 识别图片文本。
                ocr_text, ocr_runtime_error = _extract_ocr_text_from_png(page_png)
                # 记录 OCR 运行时错误（若有）。
                if ocr_runtime_error:
                    ocr_error = ocr_runtime_error
                # OCR 有有效文本则累计计数。
                if ocr_text:
                    ocr_pages += 1

        # 默认本页没有 OSS 图片地址。
        oss_url = None
        # 同时满足“本页有 PNG”和“上传开关开启”才尝试上传。
        if page_png and _env_bool("PAGE_INDEX_RAG_OSS_UPLOAD_ENABLED", True):
            # 构造稳定文件名：文档名_页码.png。
            asset_name = f"{doc_stem}_p{page_idx}.png"
            # 调用 Java 代理接口上传到 OSS。
            oss_url = _upload_image_to_java_oss(
                image_bytes=page_png,
                file_name=asset_name,
                document_name=doc_stem,
            )
            # 上传成功则记录资产信息到 metadata。
            if oss_url:
                page_image_assets.append(
                    {"page_no": page_idx, "file_name": asset_name, "oss_url": oss_url}
                )

        # 合并“原生文本 + OCR 文本 + 图片资产链接”为单页最终文本。
        merged = _build_hybrid_page_text(native_text=native_text, ocr_text=ocr_text, oss_url=oss_url)
        # 将合并结果加入输出页列表。
        result_pages.append(merged)

    # 若所有页最终都没有有效文本，则作为失败处理。
    if not any(p.strip() for p in result_pages):
        # 若有 OCR 错误，附加错误原因。
        if ocr_error:
            raise ValueError(
                "PDF contains no extractable text and OCR failed. "
                f"reason={ocr_error}"
            )
        # 无 OCR 错误信息时给出通用提示。
        raise ValueError("PDF contains no extractable text.")

    # 汇总 OCR 与图片资产元数据，供上游写入文档 metadata。
    metadata: dict[str, Any] = {
        "ocr_summary": {
            # 解析后总页数。
            "total_pages": len(result_pages),
            # 原生可提取文本页数。
            "native_text_pages": native_pages,
            # 实际采用 OCR 且识别出文本的页数。
            "ocr_pages": ocr_pages,
            # 当前流程已启用 OCR 能力。
            "ocr_enabled": True,
        },
        # 每页上传成功的图片资产列表。
        "page_image_assets": page_image_assets,
    }

    # 只要有一页用到 OCR，就标记为 pdf+ocr。
    source_type = "pdf+ocr" if ocr_pages > 0 else "pdf"
    # 返回统一解析结果。
    return ParsedDocument(pages=result_pages, source_type=source_type, metadata=metadata)


# 抽取 PDF 文本（宽松模式）：保留空页，交给上层按页决定是否 OCR。
def _extract_pdf_pages_allow_empty(data: bytes) -> list[str]:
    try:
        # 延迟导入 pypdf，减少非 PDF 路径的依赖负担。
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - import error depends on environment
        # 依赖缺失时抛出清晰提示。
        raise RuntimeError("pypdf is required for PDF upload. Please install pypdf.") from exc

    # 从内存字节流创建 PDF 读取器。
    reader = PdfReader(BytesIO(data))
    # 结果列表允许出现空字符串（表示该页无原生文本）。
    pages: list[str] = []
    # 按页提取文本并规范化后加入列表。
    for page in reader.pages:
        pages.append(normalize_text(page.extract_text() or ""))
    # 返回与 PDF 页数一一对应的列表。
    return pages


# 把指定 PDF 页渲染为 PNG 字节，供 OCR 使用。
def _render_pdf_page_to_png(data: bytes, page_index: int) -> tuple[bytes | None, str | None]:
    try:
        # PyMuPDF 导入成功才能进行渲染。
        import fitz  # type: ignore
    except Exception as exc:  # pragma: no cover - import error depends on environment
        # 返回错误字符串而非抛异常，便于上层继续流程并记录错误。
        return None, f"pymupdf_unavailable: {exc}"

    # 从环境变量读取渲染倍率（越大越清晰，但耗时更高）。
    scale = float(os.getenv("PAGE_INDEX_RAG_OCR_RENDER_SCALE", "2.0"))
    try:
        # 以内存方式打开 PDF，避免落地临时文件。
        with fitz.open(stream=data, filetype="pdf") as pdf_doc:
            # 边界检查：页索引非法时直接返回错误。
            if page_index < 0 or page_index >= len(pdf_doc):
                return None, f"page_index_out_of_range: {page_index}"
            # 获取目标页对象。
            page = pdf_doc[page_index]
            # 创建缩放矩阵。
            matrix = fitz.Matrix(scale, scale)
            # 渲染为像素图（alpha=False 减少体积）。
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            # 输出 PNG 字节并返回成功。
            return pix.tobytes("png"), None
    except Exception as exc:  # pragma: no cover - runtime depends on file quality
        # 运行时异常统一包装为 render_failed。
        return None, f"render_failed: {exc}"


# OCR 引擎单例缓存，避免每页重复初始化模型。
_OCR_ENGINE = None


# 对 PNG 字节执行 OCR，返回（识别文本, 错误信息）。
def _extract_ocr_text_from_png(image_bytes: bytes) -> tuple[str, str | None]:
    # 需要修改模块级缓存变量，所以声明 global。
    global _OCR_ENGINE
    try:
        # 延迟导入 OCR 依赖，降低冷启动成本。
        from rapidocr_onnxruntime import RapidOCR  # type: ignore
    except Exception as exc:  # pragma: no cover - import error depends on environment
        # 依赖缺失时返回空文本+错误，不中断整个文档流程。
        return "", f"rapidocr_unavailable: {exc}"

    try:
        # 首次调用时初始化 OCR 引擎。
        if _OCR_ENGINE is None:
            _OCR_ENGINE = RapidOCR()
        # 执行 OCR；第二个返回值当前未使用。
        result, _ = _OCR_ENGINE(image_bytes)
    except Exception as exc:  # pragma: no cover - runtime depends on OCR runtime
        # OCR 运行失败时返回错误描述。
        return "", f"ocr_runtime_failed: {exc}"

    # 解析 OCR 结构化结果，展开为纯文本。
    text = _flatten_ocr_result(result)
    # 规范化后返回，无错误。
    return normalize_text(text), None


# 把 OCR 多种可能结构展平为换行文本。
def _flatten_ocr_result(result: Any) -> str:
    # 空结果直接返回空字符串。
    if not result:
        return ""
    # 存放抽取到的每行文本。
    lines: list[str] = []
    # 遍历 OCR 每个元素。
    for item in result:
        # 兼容 tuple/list 结构。
        if isinstance(item, (list, tuple)):
            # 常见结构：[box, "text", score]
            if len(item) >= 2 and isinstance(item[1], str):
                lines.append(item[1])
                continue
            # 兼容嵌套结构：[box, ("text", score), ...]
            if len(item) >= 2 and isinstance(item[1], (list, tuple)) and item[1]:
                if isinstance(item[1][0], str):
                    lines.append(item[1][0])
                    continue
        # 兼容 dict 结构（某些 OCR 后处理会返回对象）。
        if isinstance(item, dict):
            # 读取 text 字段。
            text = item.get("text")
            # 只接受字符串文本。
            if isinstance(text, str):
                lines.append(text)
    # 过滤空白行后按换行拼接。
    return "\n".join(s for s in lines if s and s.strip())


# 合并单页文本：原生文本 + OCR 文本 + 图片资产 URL。
def _build_hybrid_page_text(*, native_text: str, ocr_text: str, oss_url: str | None) -> str:
    # 使用列表拼装，最后统一 join，避免手写换行分支。
    parts: list[str] = []
    # 原生文本存在则优先保留。
    if native_text:
        parts.append(native_text)
    # OCR 文本存在时再判断是否重复。
    if ocr_text:
        # 如果 OCR 文本已包含在原生文本中，则跳过，避免重复。
        if native_text and ocr_text in native_text:
            pass
        else:
            # 用标记头区分这段内容来自 OCR。
            parts.append(f"[OCR文本]\n{ocr_text}")
    # 存在图片 URL 时附加资产标记，便于下游引用展示。
    if oss_url:
        parts.append(f"[图片资产]: ossUrl={oss_url}")
    # 段落之间用空行拼接，并去掉首尾空白。
    merged = "\n\n".join(parts).strip()
    # 返回单页最终文本。
    return merged


# 调用 Java `/assets/upload-image` 接口上传图片，成功返回 oss_url。
def _upload_image_to_java_oss(*, image_bytes: bytes, file_name: str, document_name: str) -> str | None:
    # 读取上传目标地址；默认指向本地 Java 服务。
    upload_url = os.getenv(
        "PAGE_INDEX_RAG_OSS_UPLOAD_URL",
        "http://127.0.0.1:8787/api/douya/page-index-rag/assets/upload-image",
    ).strip()
    # 上传地址为空时，直接视为关闭上传能力。
    if not upload_url:
        return None

    # 构造 Java 接口所需 JSON 载荷。
    payload = {
        # 资产文件名（例如 xxx_p1.png）。
        "file_name": file_name,
        # 文档名（用于服务端拼接 OSS 目录）。
        "document_name": document_name,
        # 图片字节转 base64 文本，便于 JSON 传输。
        "content_base64": base64.b64encode(image_bytes).decode("ascii"),
    }
    # 请求超时时间，支持环境变量覆盖。
    timeout_seconds = _env_int("PAGE_INDEX_RAG_OSS_UPLOAD_TIMEOUT_SECONDS", 12)

    try:
        # 构造 POST 请求对象。
        req = Request(
            # 请求 URL。
            url=upload_url,
            # JSON 字符串转 UTF-8 字节。
            data=json.dumps(payload).encode("utf-8"),
            # 指定 JSON 请求头。
            headers={"Content-Type": "application/json"},
            # 明确 HTTP 方法为 POST。
            method="POST",
        )
        # 发起请求并等待响应（受 timeout 控制）。
        with urlopen(req, timeout=timeout_seconds) as resp:  # noqa: S310
            # 按 UTF-8 解码响应体。
            body = resp.read().decode("utf-8")
        # 解析响应 JSON。
        data = json.loads(body)
        # 仅 SUCCESS 视为成功上传。
        if str(data.get("status", "")).upper() == "SUCCESS":
            # 读取 oss_url 并去空白。
            oss_url = str(data.get("oss_url", "")).strip()
            # 空字符串时返回 None。
            return oss_url or None
        # 非 SUCCESS 记录告警日志。
        logger.warning("java_oss_upload_failed status=%s body=%s", data.get("status"), body[:300])
    except (URLError, TimeoutError, json.JSONDecodeError, OSError) as exc:
        # 请求层异常统一打告警，流程继续（不因上传失败中断入库）。
        logger.warning("java_oss_upload_request_failed file=%s error=%s", file_name, exc)
    # 上传失败统一返回 None。
    return None


# 生成安全文档 stem：仅保留字母/数字/._-，其他字符替换为下划线。
def _sanitize_file_stem(filename: str) -> str:
    # 提取不带后缀的文件名；为空时给默认值。
    stem = Path(filename).stem or "uploaded_document"
    # 逐字符过滤，避免非法路径字符进入对象名。
    return "".join(ch if (ch.isalnum() or ch in {"_", "-", "."}) else "_" for ch in stem)


# 从环境变量读取布尔值（支持常见真值表达）。
def _env_bool(name: str, default: bool) -> bool:
    # 读取原始字符串。
    raw = os.getenv(name)
    # 未设置时回落默认值。
    if raw is None:
        return default
    # 标准化后判断是否属于真值集合。
    return raw.strip().lower() in {"1", "true", "yes", "y", "on"}


# 从环境变量读取整数；非法值回落默认。
def _env_int(name: str, default: int) -> int:
    # 读取原始字符串。
    raw = os.getenv(name)
    # 未设置时直接返回默认值。
    if raw is None:
        return default
    try:
        # 尝试解析整数。
        return int(raw)
    except ValueError:
        # 解析失败回落默认值，避免抛异常中断流程。
        return default


# 解析 DOCX 文本：提取所有非空段落并用空行拼接。
def _parse_docx_text(data: bytes) -> str:
    try:
        # 运行时导入 python-docx。
        from docx import Document
    except Exception as exc:  # pragma: no cover - import error depends on environment
        # 依赖缺失时给安装提示。
        raise RuntimeError("python-docx is required for DOCX upload. Please install python-docx.") from exc

    # 从内存字节流创建 docx 文档对象。
    doc = Document(BytesIO(data))
    # 收集清洗后的段落文本。
    lines: list[str] = []
    # 遍历所有段落。
    for p in doc.paragraphs:
        # 规范化段落文本。
        text = normalize_text(p.text)
        # 非空段落才保留。
        if text:
            lines.append(text)
    # 没有任何有效文本时抛错。
    if not lines:
        raise ValueError("DOCX contains no extractable text.")
    # 段落之间用空行拼接，提升可读性。
    return "\n\n".join(lines)
